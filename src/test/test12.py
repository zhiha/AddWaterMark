import os
from win32com.client import DispatchEx
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2 import PdfFileWriter, PdfFileReader
import pandas as pd
from docx import Document
from docx.shared import Pt  #设置像素、缩进等
from docx.shared import RGBColor #设置字体颜色
from docx.oxml.ns import qn
from hashlib import md5
from PyPDF4 import PdfFileReader, PdfFileWriter
from PyPDF4.generic import NameObject, DictionaryObject, ArrayObject, NumberObject, ByteStringObject
from PyPDF4.pdf import _alg33, _alg34, _alg35
from PyPDF4.utils import b_
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTText, LTChar, LTAnno
import shutil
import glob, fitz
from pathlib import Path
import img2pdf

pdfmetrics.registerFont(TTFont('kaiti', 'C:/Windows/Fonts/simkai.ttf'))  # 楷体


TEMP_DIR = os.path.join(os.getcwd(), 'temp')
MINUTE_DIR = os.path.join(os.getcwd(), 'result')
INPUT_DIR = os.path.join(os.getcwd(), 'inputs')
if not os.path.exists(TEMP_DIR):
    os.mkdir(TEMP_DIR)
if not os.path.exists(MINUTE_DIR):
    os.mkdir(MINUTE_DIR)
if not os.path.exists(INPUT_DIR):
    os.mkdir(INPUT_DIR)

word_abs_path = os.path.join(os.getcwd(), 'inputs', [x for x in os.listdir(os.path.join(os.getcwd(), 'inputs')) if
                                                     x.endswith('.docx') or x.endswith('.doc')][0])
excel_abs_path = os.path.join(os.getcwd(), 'inputs', [x for x in os.listdir(os.path.join(os.getcwd(), 'inputs')) if
                                                      x.endswith('.xlsx') or x.endswith('.xls')][0])


def convert_word2pdf(word_abs_path):
    app = DispatchEx('Word.Application')
    app.Visible = 0  # 这个至少在调试阶段建议打开，否则如果等待时间长的话，它至少给你耐心。。。
    app.DisplayAlerts = 0
    doc = app.Documents.Open(word_abs_path)

    all_content = doc.Range(doc.Content.Start, doc.Content.End)
    # all_content.HighlightColorIndex = 16  # 全局背景色淡黄色

    temp_pdf_abs_path = os.path.join(TEMP_DIR,
                                     os.path.basename(word_abs_path).replace('.docx', '.pdf').replace('.doc', '.pdf'))

    doc.SaveAs(temp_pdf_abs_path, FileFormat=17)
    doc.Close()
    app.Quit()
    return temp_pdf_abs_path  # 返回临时pdf的路径


####### 1.生成水印pdf的函数 ########
def create_watermark(content,target_path):
    fp = open(target_path, 'rb')
    parser = PDFParser(fp)
    doc: PDFDocument = PDFDocument(parser)
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    space = 0.825
    pre_process = 0
    value_est1 = 0
    value_est2 = 0
    space_raw = 0
    space_factor = 0

    page_num = 0
    for page in PDFPage.create_pages(doc):
        c = canvas.Canvas(os.path.join(TEMP_DIR, 'watermark-%i.pdf') % page_num , pagesize=(21 * cm, 29.7 * cm))
        c.setFillColorRGB(190 / 255, 190 / 255, 190 / 255, alpha=0.4)  # 淡
        c.translate(0 * cm, 27 * cm)  # 移动坐标原点(坐标系左下为(0,0)))
        c.setFont('kaiti', 10.5)
        for i in [3, 8.5, 14]:
            for j in range(0, 32, 2):
                c.drawString(i * cm, -space * j * cm, content)
        page_num = page_num + 1
        interpreter.process_page(page)
        layout = device.get_result()
        pre = 800
        cur = 0
        cnt = 0
        for textbox in layout:
            if pre_process == 2:
                break
            if isinstance(textbox, LTText):
                for line in textbox:
                    if pre_process < 2:
                        if pre_process == 0:
                            value_est1 = line.bbox[3]
                        if pre_process == 1:
                            value_est2 = line.bbox[3]
                            space_raw = value_est1 - value_est2
                            space_factor = space_raw / space
                        pre_process = pre_process + 1
                    else: break
        for textbox in layout:
            if isinstance(textbox, LTText):
                for line in textbox:
                    print("line site:", line.bbox[0], "y:", line.bbox[3], line.width, line.height)
                    if line.width < 300: # and (pre - cur > 22):
                        if cnt % 4 == 0 :
                            c.setFillColorRGB(190 / 255, 190 / 255, 190 / 255, alpha=1)  # 淡
                            c.setFont('kaiti', 6)
                            c.drawString(15*cm, ((line.bbox[3] - value_est1) / space_factor - space/2) * cm, content)
                        cnt = cnt + 1
                    # pre = cur
        c.save()  # 关闭并保存pdf文件
    fp.close()

######## 2.为pdf文件加水印的函数 ########
def add_watermark2pdf(input_pdf, output_pdf, TEMP_DIR):
    pdf = PdfFileReader(input_pdf, strict=False)
    pdf_writer = PdfFileWriter()
    page_num = 0
    for page in range(pdf.getNumPages()):
        watermark_pdf = os.path.join(TEMP_DIR, 'watermark-%i.pdf') % page_num
        page_num = page_num + 1
        watermark = PdfFileReader(watermark_pdf)
        watermark_page = watermark.getPage(0)
        pdf_page = pdf.getPage(page)
        pdf_page.mergePage(watermark_page)
        pdf_page.compressContentStreams()
        pdf_writer.addPage(pdf_page)
    pdfOutputFile = open(output_pdf, 'wb')
    pdf_writer.write(pdfOutputFile)
    pdfOutputFile.close()

def formatWord(input_word_path,output_word_path,space):
    doc = Document(input_word_path)
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = Pt(space)#1.5
        for run in paragraph.runs:
            run.font.bold = True
            run.font.italic = False
            run.font.underline = False
            run.font.strike = False
            run.font.shadow = False
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = "等线"

            # 设置像黑体这样的中文字体，必须添加下面 2 行代码
            r = run._element.rPr.rFonts
            r.set(qn("w:eastAsia"), "等线")
    doc.save(output_word_path)

def encrypt(writer_obj: PdfFileWriter, user_pwd, owner_pwd=None, use_128bit=True):
    import time, random
    if owner_pwd == None:
        owner_pwd = user_pwd
    if use_128bit:
        V = 2
        rev = 3
        keylen = int(128 / 8)
    else:
        V = 1
        rev = 2
        keylen = int(40 / 8)
    # permit copy and printing only:
    P = -3904
    O = ByteStringObject(_alg33(owner_pwd, user_pwd, rev, keylen))
    ID_1 = ByteStringObject(md5(b_(repr(time.time()))).digest())
    ID_2 = ByteStringObject(md5(b_(repr(random.random()))).digest())
    writer_obj._ID = ArrayObject((ID_1, ID_2))
    if rev == 2:
        U, key = _alg34(user_pwd, O, P, ID_1)
    else:
        assert rev == 3
        U, key = _alg35(user_pwd, rev, keylen, O, P, ID_1, False)
    encrypt = DictionaryObject()
    encrypt[NameObject("/Filter")] = NameObject("/Standard")
    encrypt[NameObject("/V")] = NumberObject(V)
    if V == 2:
        encrypt[NameObject("/Length")] = NumberObject(keylen * 8)
    encrypt[NameObject("/R")] = NumberObject(rev)
    encrypt[NameObject("/O")] = ByteStringObject(O)
    encrypt[NameObject("/U")] = ByteStringObject(U)
    encrypt[NameObject("/P")] = NumberObject(P)
    writer_obj._encrypt = writer_obj._addObject(encrypt)
    writer_obj._encrypt_key = key

# def parse_char_layout(layout):
#     """解析页面内容，一个字母一个字母的解析"""
#     # bbox:
#     # x0：从页面左侧到框左边缘的距离。
#     # y0：从页面底部到框的下边缘的距离。
#     # x1：从页面左侧到方框右边缘的距离。
#     # y1：从页面底部到框的上边缘的距离
#     for textbox in layout:
#         if isinstance(textbox, LTText):
#             for line in textbox:
#                 print("line site:", line.bbox[0], "y:", line.bbox[3], line.width , line.height)

def main():
    print('程序正在运行…………by Superon')
    ### temp路径
    temp_word_abs_path = os.path.join(TEMP_DIR,os.path.basename(word_abs_path))

    ### 格式化word
    # formatWord(word_abs_path,temp_word_abs_path,18)

    ### word转pdf
    temp_pdf_abs_path = convert_word2pdf(word_abs_path)

    ### 获取人员名单
    persons = pd.read_excel(excel_abs_path).to_dict('records')

    for person in persons:
        ### 创建水印PDF
        wtmk_content = '仅供%s-%s参考' % (person['fund_company'], person['reseacher'])
        print(wtmk_content)
        # create_watermark(wtmk_content)
        ### 合并纪要pdf和水印pdf
        # watermark_pdf = os.path.join(TEMP_DIR, 'watermark.pdf')
        input_pdf = temp_pdf_abs_path
        output_pdf = os.path.join(TEMP_DIR, os.path.splitext(os.path.basename(word_abs_path))[0] + '_' + person[
            'fund_company'] + '_' + person['reseacher'] + "_tmp" + '.pdf')
        create_watermark(wtmk_content,input_pdf)
        add_watermark2pdf(input_pdf, output_pdf, TEMP_DIR)

        # To get better resolution
        zoom_x = 4.0  # horizontal zoom
        zoom_y = 4.0  # vertical zoom
        mat = fitz.Matrix(zoom_x, zoom_y)  # zoom factor 2 in each dimension
        doc = fitz.open(output_pdf)  # open document
        for page in doc:  # iterate through the pages
            pix = page.get_pixmap(matrix=mat)  # render page to an image
            pix.save(os.path.join(TEMP_DIR,"page-%i.png") % page.number)  # store image as a PNG
        doc.close()

        output_pdf = os.path.join(TEMP_DIR, os.path.splitext(os.path.basename(word_abs_path))[0] + '_' + person[
            'fund_company'] + '_' + person['reseacher'] + '.pdf')
        with open(output_pdf, "wb") as f:
            f.write(img2pdf.convert([str(path) for path in Path(TEMP_DIR).glob('*.png')]))
            f.close()

        ### 加权限
        unmeta = PdfFileReader(output_pdf,strict=False)
        writer = PdfFileWriter()
        writer.appendPagesFromReader(unmeta)
        encrypt(writer, '', '123')
        final_pdf = os.path.join(MINUTE_DIR, os.path.splitext(os.path.basename(word_abs_path))[0] + '_' + person[
            'fund_company'] + '_' + person['reseacher'] + '.pdf')
        with open(final_pdf, 'wb') as fp:
            writer.write(fp)

    shutil.rmtree(TEMP_DIR)


if __name__ == '__main__':
    main()